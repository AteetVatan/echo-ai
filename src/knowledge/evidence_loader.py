"""
Evidence Vault loader (Layer B).

Loads supporting evidence documents from ``rag_persona_db/document/``:
- GitHub READMEs (.md) — structure-aware chunking by headers
- CV (.docx) — paragraph-level chunking
- LinkedIn CSV exports — selective parsing of high-value files
- PDFs / plain text — standard chunking

Each chunk receives metadata for retrieval filtering and traceability.
"""

from __future__ import annotations

import csv
import hashlib
import io
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import (
    MarkdownHeaderTextSplitter,
    RecursiveCharacterTextSplitter,
)

from src.utils import get_logger

logger = get_logger(__name__)

# LinkedIn data files worth ingesting (rest is low-value or private)
_LINKEDIN_HIGH_VALUE_FILES: set[str] = {
    "profile.csv",
    "positions.csv",
    "education.csv",
    "skills.csv",
    "certifications.csv",
    "projects.csv",
    "endorsement_received_info.csv",
    "recommendations_given.csv",
    "learning.csv",
    "languages.csv",
}

# Chunk sizes
_MD_CHUNK_SIZE = 1000
_MD_CHUNK_OVERLAP = 150
_DEFAULT_CHUNK_SIZE = 800
_DEFAULT_CHUNK_OVERLAP = 100


def _make_evidence_stable_id(source: str, chunk_index: int) -> str:
    """Deterministic ID for an evidence chunk."""
    raw = f"{source}:chunk:{chunk_index}"
    return hashlib.sha256(raw.encode()).hexdigest()[:16]


# ------------------------------------------------------------------
# Markdown (GitHub READMEs)
# ------------------------------------------------------------------

def _load_markdown(path: Path) -> list[Document]:
    """Split markdown by headers, then sub-split large sections."""
    text = path.read_text(encoding="utf-8", errors="replace")

    header_splitter = MarkdownHeaderTextSplitter(
        headers_to_split_on=[
            ("#", "h1"),
            ("##", "h2"),
            ("###", "h3"),
        ],
        strip_headers=False,
    )
    header_docs = header_splitter.split_text(text)

    # Sub-split large sections
    sub_splitter = RecursiveCharacterTextSplitter(
        chunk_size=_MD_CHUNK_SIZE,
        chunk_overlap=_MD_CHUNK_OVERLAP,
    )

    chunks: list[Document] = []
    for doc in header_docs:
        sub_chunks = sub_splitter.split_text(doc.page_content)
        for sc in sub_chunks:
            chunks.append(
                Document(
                    page_content=sc,
                    metadata={
                        **doc.metadata,
                        "source": path.name,
                        "parent_doc_id": path.stem,
                    },
                )
            )
    return chunks


# ------------------------------------------------------------------
# DOCX (CV)
# ------------------------------------------------------------------

def _load_docx(path: Path) -> list[Document]:
    """Load .docx and split by paragraphs / recursive chunks."""
    try:
        from docx import Document as DocxDocument  # python-docx
    except ImportError:
        logger.warning("python-docx not installed, skipping %s", path.name)
        return []

    doc = DocxDocument(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=_DEFAULT_CHUNK_SIZE,
        chunk_overlap=_DEFAULT_CHUNK_OVERLAP,
    )
    chunks = splitter.split_text(full_text)
    return [
        Document(
            page_content=chunk,
            metadata={"source": path.name, "parent_doc_id": path.stem},
        )
        for chunk in chunks
    ]


# ------------------------------------------------------------------
# CSV (LinkedIn data exports)
# ------------------------------------------------------------------

def _load_csv(path: Path) -> list[Document]:
    """Parse a LinkedIn CSV into one document per row (or combined)."""
    text = path.read_text(encoding="utf-8", errors="replace")
    reader = csv.DictReader(io.StringIO(text))

    rows: list[str] = []
    for row in reader:
        # Build readable text from each row
        parts = [f"{k}: {v}" for k, v in row.items() if v and v.strip()]
        if parts:
            rows.append(" | ".join(parts))

    if not rows:
        return []

    # Combine rows into chunks of reasonable size
    combined = "\n".join(rows)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=_DEFAULT_CHUNK_SIZE,
        chunk_overlap=_DEFAULT_CHUNK_OVERLAP,
    )
    chunks = splitter.split_text(combined)
    return [
        Document(
            page_content=chunk,
            metadata={"source": path.name, "parent_doc_id": path.stem},
        )
        for chunk in chunks
    ]


# ------------------------------------------------------------------
# Plain text / PDF fallback
# ------------------------------------------------------------------

def _load_text(path: Path) -> list[Document]:
    """Fallback loader for .txt files."""
    text = path.read_text(encoding="utf-8", errors="replace")
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=_DEFAULT_CHUNK_SIZE,
        chunk_overlap=_DEFAULT_CHUNK_OVERLAP,
    )
    chunks = splitter.split_text(text)
    return [
        Document(
            page_content=chunk,
            metadata={"source": path.name, "parent_doc_id": path.stem},
        )
        for chunk in chunks
    ]


def _load_pdf(path: Path) -> list[Document]:
    """Load PDF if PyPDF is available."""
    try:
        from langchain_community.document_loaders import PyPDFLoader
    except ImportError:
        logger.warning("pypdf not installed, skipping %s", path.name)
        return []

    loader = PyPDFLoader(str(path))
    pages = loader.load()

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=_DEFAULT_CHUNK_SIZE,
        chunk_overlap=_DEFAULT_CHUNK_OVERLAP,
    )
    chunks: list[Document] = []
    for page in pages:
        sub_chunks = splitter.split_text(page.page_content)
        for sc in sub_chunks:
            chunks.append(
                Document(
                    page_content=sc,
                    metadata={
                        "source": path.name,
                        "parent_doc_id": path.stem,
                        "page": page.metadata.get("page", 0),
                    },
                )
            )
    return chunks


# ------------------------------------------------------------------
# Public API
# ------------------------------------------------------------------

def load_evidence_documents(evidence_dir: Path | str) -> list[Document]:
    """Load all evidence documents from *evidence_dir* and its subdirectories.

    Parameters
    ----------
    evidence_dir:
        Root directory (typically ``rag_persona_db/document/``).

    Returns
    -------
    list[Document]
        Chunked documents with ``layer="evidence"`` metadata and stable IDs.
    """
    evidence_dir = Path(evidence_dir)
    if not evidence_dir.exists():
        logger.warning("Evidence directory does not exist: %s", evidence_dir)
        return []

    all_docs: list[Document] = []

    for path in sorted(evidence_dir.rglob("*")):
        if not path.is_file():
            continue

        suffix = path.suffix.lower()
        name_lower = path.name.lower()

        # Filter LinkedIn CSVs
        if suffix == ".csv":
            if name_lower not in _LINKEDIN_HIGH_VALUE_FILES:
                logger.debug("Skipping low-value CSV: %s", path.name)
                continue
            chunks = _load_csv(path)
        elif suffix == ".md":
            chunks = _load_markdown(path)
        elif suffix == ".docx":
            chunks = _load_docx(path)
        elif suffix == ".pdf":
            chunks = _load_pdf(path)
        elif suffix == ".txt":
            chunks = _load_text(path)
        else:
            logger.debug("Skipping unsupported file: %s", path.name)
            continue

        if chunks:
            logger.info("Loaded %d chunks from %s", len(chunks), path.name)

        all_docs.extend(chunks)

    # Assign stable IDs and layer metadata
    for idx, doc in enumerate(all_docs):
        source = doc.metadata.get("source", "unknown")
        doc.metadata["stable_id"] = _make_evidence_stable_id(source, idx)
        doc.metadata["layer"] = "evidence"

    logger.info("Total evidence chunks loaded: %d", len(all_docs))
    return all_docs
